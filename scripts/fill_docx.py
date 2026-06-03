"""Fill in the 2026 Open Lab Student Research Project Application DOCX."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOCX_PATH = "docs/2026年开放实验室学生科研项目申请书-已填写.docx"

doc = Document(DOCX_PATH)

# ── Font helper ──────────────────────────────────────────────
def set_fangsong(run, size=Pt(12), bold=False):
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = size
    run.bold = bold


def add_fangsong_paragraph(cell, text, size=Pt(12), bold=False, alignment=None, first_line_indent=None):
    para = cell.add_paragraph()
    run = para.add_run(text)
    set_fangsong(run, size, bold)
    if alignment is not None:
        para.alignment = alignment
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent
    para.paragraph_format.line_spacing = 1.5
    return para


# ═══════════════════════════════════════════════════════════════
# TABLE 0: Project Basic Info
# ═══════════════════════════════════════════════════════════════
t0 = doc.tables[0]

# ── Project Name (Row 17, Cell 1) ──
name_cell = t0.cell(17, 1)
name_cell.text = ""
add_fangsong_paragraph(
    name_cell,
    "基于大语言模型的AI Native自然语言事务管理系统的设计与实现",
    size=Pt(14), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
)

# ── Project Type (Row 18, Cell 1) ──
type_cell = t0.cell(18, 1)
type_cell.text = ""
add_fangsong_paragraph(
    type_cell,
    "☑ 一般项目（耗材经费1000元）  □ 重点项目（耗材经费1500元）",
    size=Pt(12)
)

# ── Deliverables (Row 19, Cell 1) ──
result_cell = t0.cell(19, 1)
result_cell.text = ""
add_fangsong_paragraph(
    result_cell,
    "☑ 软件著作权  □ 论文  □ 调研报告  □ 专利  □ 科技孵化产品  □ 模型设计  □ 艺术设计  □ 艺术作品等",
    size=Pt(12)
)


# ═══════════════════════════════════════════════════════════════
# TABLE 1: Project Details + Budget
# ═══════════════════════════════════════════════════════════════
t1 = doc.tables[1]

# ── Content Cell (Row 0, Cell 0, span=4) ──
content_cell = t1.cell(0, 0)
for p in content_cell.paragraphs:
    p.clear()

# Remove all paragraphs except one, then remove that one
while len(content_cell.paragraphs) > 1:
    p = content_cell.paragraphs[-1]
    p._element.getparent().remove(p._element)
first_p = content_cell.paragraphs[0]
first_p._element.getparent().remove(first_p._element)

# ── Section 1 ──
add_fangsong_paragraph(content_cell, "1．研究概述和目的", size=Pt(12), bold=True)

sec1 = (
    "随着大语言模型（LLM）技术的快速发展，以DeepSeek、通义千问等为代表的国产大模型"
    "已具备强大的自然语言理解与结构化信息提取能力。然而，当前主流个人事务管理软件（如记账、"
    "日程、提醒类APP）仍普遍采用传统的「表单+点击」交互模式，要求用户逐步选择分类、填写金额、"
    "设定时间，操作路径冗长，学习成本高，且无法直接理解诸如「今天午饭35元，明天下午3点开会」"
    "的自然语言表达。这种「人适应软件」的设计范式已明显落后于AI时代的技术能力。"
    "本研究旨在设计并实现一款AI Native自然语言事务管理系统，核心理念为「让软件理解用户，"
    "而非让用户适应软件」。用户仅需以日常自然语言（文字或语音）输入事务需求，系统即可自动"
    "完成意图识别、信息提取、分类打标、时间解析与数据持久化等全流程操作，实现「一句话完成」"
    "记账、日程创建、提醒设置等事务管理操作。本项目的实施将为LLM技术在个人事务管理领域的"
    "工程化落地提供一套完整且可复用的技术方案，同时探索AI Native应用区别于传统软件的全新"
    "交互范式与系统架构，具有重要的实践意义与学术参考价值。"
)
add_fangsong_paragraph(content_cell, sec1, size=Pt(12), first_line_indent=Cm(0.74))

# blank line
add_fangsong_paragraph(content_cell, "", size=Pt(12))

# ── Section 2 ──
add_fangsong_paragraph(content_cell, "2．本项目研究思路、研究预期成果及研究内容", size=Pt(12), bold=True)

sec2 = (
    "研究思路：本项目采用「单LLM + 模块化Prompt Workflow」的核心架构。不同于传统的多Agent"
    "系统（多个模型互相通信），本项目使用单一DeepSeek大语言模型，配合7个职责明确的专职Prompt"
    "模块（意图检测、支出解析、收入解析、日程解析、修改解析、删除解析、反馈生成），构成完整的"
    "AI事务处理流水线。用户输入自然语言后，系统首先通过意图检测Prompt判断操作类型，随后调用"
    "对应的结构化提取Prompt生成JSON数据，再经由服务端多层校验（Markdown清洗→JSON解析→"
    "Schema验证→字段补全）确保数据合法性，最终完成数据持久化并返回自然语言反馈。同时，"
    "采用LLM语义理解与Python规则引擎相结合的混合式时间解析策略，以消除大模型在时间推理"
    "方面的幻觉问题。前后端采用Flask + React技术栈，实现响应式Web应用。"
    "研究内容主要包括：（1）面向事务管理场景的模块化Prompt工程设计与迭代优化；（2）LLM输出"
    "的JSON结构化稳定性保障机制，包括多层校验与自动修复策略；（3）自然语言时间的混合解析引擎"
    "设计；（4）基于Flask的后端API与React前端SPA的系统集成；（5）多意图事务（记账+日程）的"
    "自动拆分与并发处理；（6）AI Native交互界面设计，包括AI核心球动画反馈系统。"
    "预期成果：（1）完成一款功能完整的AI Native事务管理Web应用原型，支持AI记账、AI日程、"
    "AI查询、AI修改与删除等核心功能；（2）形成一套可复用的模块化Prompt模板库；（3）申请"
    "软件著作权1项；（4）撰写技术研究报告或学术论文1篇。"
)
add_fangsong_paragraph(content_cell, sec2, size=Pt(12), first_line_indent=Cm(0.74))

# blank line
add_fangsong_paragraph(content_cell, "", size=Pt(12))

# ── Section 3 ──
add_fangsong_paragraph(content_cell, "3．研究项目的创新点", size=Pt(12), bold=True)

sec3 = (
    "（1）AI Native交互范式创新：颠覆传统事务管理APP「表单驱动」的交互模式，提出「自然语言驱动」"
    "的全新范式。用户无需学习操作流程，系统自动理解意图并完成事务操作，真正实现从「人适应软件」"
    "到「软件理解人」的范式转变。"
    "（2）模块化Prompt Workflow架构创新：不同于常见的单一巨型Prompt或多Agent框架，本项目设计"
    "「单LLM + 7个专职Prompt模块」的轻量级流水线架构，兼具多Agent职责分离优势（独立调优测试、"
    "灵活扩展），同时避免多模型调度的复杂性与高成本，在实用性与可维护性间取得良好平衡。"
    "（3）混合式时间解析机制创新：将LLM语义理解与规则引擎确定性计算相结合，由LLM理解「明天下午」"
    "「下周三」等模糊时间表达，规则引擎负责具体日期计算与标准化，有效解决纯LLM方案在时间推理中"
    "的幻觉与不一致问题。"
    "（4）JSON稳定性保障机制创新：针对LLM输出不稳定的痛点，设计Markdown清洗、JSON修复、Schema"
    "校验、默认值补全四层校验流水线，结合temperature=0确定性推理策略，显著提升事务场景结构化"
    "输出可靠性。"
)
add_fangsong_paragraph(content_cell, sec3, size=Pt(12), first_line_indent=Cm(0.74))

# ── Budget Summary (Row 1, Cell 0) ──
budget_cell = t1.cell(1, 0)
budget_cell.text = ""
add_fangsong_paragraph(budget_cell, "5．经费概算：                   共计 1000 元", size=Pt(12))

# ── Budget Table (Rows 3-5) ──
budget_data = [
    (3, ["1", "API调用费", "DeepSeek/通义千问等大语言模型API Token采购，用于系统开发调试、Prompt调优测试及演示期间的AI推理调用", "700"]),
    (4, ["2", "云计算资源", "阿里云/腾讯云ECS云服务器（按量付费），用于系统后端API部署、前端Web应用托管及公网演示环境搭建", "250"]),
    (5, ["3", "域名与SSL证书", "测试域名注册及SSL安全证书，用于Web应用公网访问与HTTPS安全连接", "50"]),
]

for row_idx, values in budget_data:
    for col_idx, text in enumerate(values):
        cell = t1.cell(row_idx, col_idx)
        cell.text = ""
        align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
        add_fangsong_paragraph(cell, text, size=Pt(12), alignment=align)

# Clear unused budget rows (rows 6-7)
for row_idx in [6, 7]:
    for col_idx in range(4):
        t1.cell(row_idx, col_idx).text = ""

# ── Save ──
doc.save(DOCX_PATH)
print(f"Saved to {DOCX_PATH}")
print("Done!")
